import json
import re
import openpyxl
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

# -----------------------------
# Setup
# -----------------------------
load_dotenv()
client = OpenAI(api_key=None)  # Set API key in environment or Streamlit

# Token trackers
total_input_tokens = 0
total_output_tokens = 0

# Pricing constants
PRICE_INPUT = 0.00000015
PRICE_OUTPUT = 0.00000060
USD_TO_RM = 4.14

# -----------------------------
# Helper function to update tokens
# -----------------------------
def update_token_usage(usage):
    """Update global token counters for every API call."""
    global total_input_tokens, total_output_tokens
    total_input_tokens += getattr(usage, "prompt_tokens", 0)
    total_output_tokens += getattr(usage, "completion_tokens", 0)

# -----------------------------
# 1. Load Excel as text
# -----------------------------
def load_excel_as_text(file):
    wb = openpyxl.load_workbook(file, data_only=True)
    sheet_texts = {}
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        lines = []
        for row in sheet.iter_rows(values_only=True):
            line_cells = [str(x) for x in row if x is not None]
            if line_cells:
                lines.append(" | ".join(line_cells))
        sheet_texts[sheet_name] = "\n".join(lines)
    return sheet_texts

# -----------------------------
# 2. Clean JSON output
# -----------------------------
def clean_json_output(raw_output):
    return re.sub(r"```(?:json)?\n(.*?)```", r"\1", raw_output, flags=re.DOTALL).strip()

# -----------------------------
# 3. Detect references with AI
# -----------------------------
def detect_references_with_ai(all_sheets_text):
    text = "\n\n".join([f"=== Sheet: {k} ===\n{v}" for k, v in all_sheets_text.items()])
    prompt = f"""
Detect references to other sheets/tabs in the following Excel text.
Output ONLY a JSON array: [{{"question_text": "...", "referenced_sheet": "..."}}]

Text:
{text}
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    update_token_usage(response.usage)
    raw_output = response.choices[0].message.content
    try:
        return json.loads(clean_json_output(raw_output))
    except:
        return []

# -----------------------------
# 4. Answer references
# -----------------------------
def answer_references(all_sheets_text, references):
    answers = []
    for ref in references:
        question = ref["question_text"]
        sheet_name = ref["referenced_sheet"]
        sheet_text = all_sheets_text.get(sheet_name)
        if not sheet_text:
            answer_text = f"[Error: Referenced sheet '{sheet_name}' not found]"
        else:
            prompt = f"""
Use ONLY the referenced sheet text to answer the question.

QUESTION:
{question}

REFERENCED SHEET ({sheet_name}):
{sheet_text}

Provide a concise, clear answer. Do NOT add explanations.
"""
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0
            )
            update_token_usage(response.usage)
            answer_text = response.choices[0].message.content.strip()
        answers.append({
            "question": question,
            "referenced_sheet": sheet_name,
            "answer": answer_text
        })
    return answers

# -----------------------------
# 5. Merge answers into text
# -----------------------------
def merge_answers_into_text(all_sheets_text, answers):
    combined_text = "\n\n".join([f"=== Sheet: {k} ===\n{v}" for k, v in all_sheets_text.items()])
    for item in answers:
        question = item["question"].strip()
        answer = item["answer"].strip()
        if not answer or answer.lower() in ["n/a", "no data"]:
            continue
        core_phrase = question.split("Pls")[0].split(" -")[-1].strip()
        pattern = rf".*{re.escape(core_phrase)}.*"
        replacement = f"- {core_phrase}: {answer}"
        combined_text = re.sub(pattern, replacement, combined_text, count=1)
    return combined_text

# -----------------------------
# 6. Generate sustainability report
# -----------------------------
def generate_sustainability_report(final_text):
    prompt = f"""
You are a sustainability reporting expert. 
From the data provided below: 
- Exclude all entries marked as "N/A", "No data", "Not available", or blank. 
- Only include sections that are supported by the data. 
- Suggested headings: Economic, Environment, Energy Management, Emissions Management, Governance, Social
- Maximum 3 sections
- Merge related data points into well-structured narrative paragraphs. 
- Keep a professional, formal sustainability-report tone. 

Extracted content:
------------------------------
{final_text}
------------------------------
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a sustainability reporting expert."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2
    )
    update_token_usage(response.usage)
    return response.choices[0].message.content.strip()

# -----------------------------
# 7. Save report to Word + print token & cost summary
# -----------------------------
def save_report_to_word(report_text, file_path):
    doc = Document()
    doc.add_heading("Sustainability Report", 0)

    for paragraph in report_text.split("\n"):
        line = paragraph.strip()

        # Convert markdown headings to Word headings
        if line.startswith("### "):  
            doc.add_heading(line.replace("### ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=3)
        else:
            doc.add_paragraph(line)

    doc.save(file_path)

    # Compute costs...
    cost_input = total_input_tokens * PRICE_INPUT
    cost_output = total_output_tokens * PRICE_OUTPUT
    total_cost_usd = cost_input + cost_output
    total_cost_rm = total_cost_usd * USD_TO_RM

    print("\n================ TOKEN & COST SUMMARY FOR SUSTAINABILITY ================")
    print(f"Total input tokens:  {total_input_tokens:,}")
    print(f"Total output tokens: {total_output_tokens:,}")
    print(f"Cost for input (USD):   ${cost_input:.8f}")
    print(f"Cost for output (USD):  ${cost_output:.8f}")
    print(f"Total cost (USD):       ${total_cost_usd:.8f}")
    print("---------------------------------------------------------------------------")
    print(f"Estimated Total cost (RM):   RM {total_cost_rm:.4f}")
    print("===========================================================================\n")

    return file_path

