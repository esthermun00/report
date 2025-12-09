# full_pipeline_memory.py
import os
import json
import pandas as pd
import openpyxl
import matplotlib.pyplot as plt
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO

# ------------------------
# SETUP
# ------------------------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

total_input_tokens = 0
total_output_tokens = 0



# 1. Load Excel into raw text
# ------------------------
def load_excel_clean(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    sheets_text = {}
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        lines = []
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join("" if v is None else str(v) for v in row)
            if row_text.strip():
                lines.append(row_text)
        sheets_text[sheet_name] = "\n".join(lines)
    return sheets_text

# ------------------------
# 2. Generate Business & Operational Overview
# ------------------------
def generate_business_overview(mdna_text, consol_text):
    combined_text = ""
    for name, text in mdna_text.items():
        combined_text += f"\n--- Sheet: {name} (MD&A) ---\n{text}\n"
    for name, text in consol_text.items():
        combined_text += f"\n--- Sheet: {name} (Consol) ---\n{text}\n"

    prompt = f"""
You are a corporate reporting expert.

Using the extracted data provided below, generate a **Business and Operational Overview** section 
of an annual report.  

--- EXAMPLE STYLE TO FOLLOW (DO NOT COPY TEXT) ---
BUSINESS AND OPERATIONAL OVERVIEW
[2-3Paragraphs describing company, strategy, growth, etc.] (Not more than 250 words)

As an overview, our business segments are as follows:
Logistics segment : [1-line description]
Warehousing segment : [1-line description]
Trading segment : [1-line description]
Others : [1-line description]
--- END OF EXAMPLE STYLE ---

Data:
{combined_text}
"""
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1500,
        temperature=0.4
    )
    update_token_usage(response.usage)
    text = response.choices[0].message.content
    return text  # keep in memory

# ------------------------
# 3. Load consol.xlsx to JSON
# ------------------------
import re

def load_consol_to_json(excel_file="consol.xlsx", sheet_name="GRP IS"):
    df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)

    header_row = None

    date_pattern = re.compile(
        r"(\d{1,2}[\s\-\/]*(jan|january|feb|february|mar|march|apr|april|may|jun|june|"
        r"jul|july|aug|august|sep|sept|september|oct|october|nov|november|dec|december)"
        r"[\s\-\/]*\d{2,4})",
        re.IGNORECASE
    )

    for i in range(len(df)):
        row = df.iloc[i].astype(str)

        # Count how many cells look like dates
        matches = [cell for cell in row if date_pattern.search(cell)]
        if len(matches) >= 2:   # header likely has at least 2 periods
            header_row = i
            break

    if header_row is None:
        raise Exception("Header row with date columns not found.")

    # Continue your normal processing
    table = df.iloc[header_row:].reset_index(drop=True)
    table.columns = table.iloc[0]
    table = table[1:]

    table = table.dropna(axis=1, how="all").dropna(axis=0, how="all")
    table.columns = [str(col).strip() for col in table.columns]

    # Fix duplicate columns
    def fix_duplicate_columns(cols):
        seen = {}
        new_cols = []
        for col in cols:
            if col not in seen:
                seen[col] = 0
                new_cols.append(col)
            else:
                seen[col] += 1
                new_cols.append(f"{col}_{seen[col]}")
        return new_cols

    table.columns = fix_duplicate_columns(table.columns)
    table = table.rename(columns={table.columns[0]: "Item"})

    cols_to_drop = [col for col in table.columns if "nan" in str(col).lower() and col != "Item"]
    table = table.drop(columns=cols_to_drop)

    table["Item"] = table["Item"].ffill()

    def safe_num(x):
        if pd.isna(x):
            return 0
        if isinstance(x, str):
            t = x.strip()
            if t in ["", "-", "–"]:
                return 0
            t = t.replace(",", "").replace("(", "-").replace(")", "")
            try:
                return float(t)
            except:
                return x
        try:
            return float(x)
        except:
            return x

    for col in table.columns:
        if col != "Item":
            table[col] = table[col].apply(safe_num)

    return table.to_dict(orient="records")


# ------------------------
# 4. Generate financial JSON via GPT
# ------------------------
def ask_ai_for_financials(raw_text_dict, clean_json):
    combined_text = "\n\n".join(f"{name}:\n{text}" for name, text in raw_text_dict.items())

    prompt = """
You are a financial extraction engine. Your job is to extract a FULL SET of financial metrics for ALL YEARS that appear
in the data.

You MUST extract the following for every year:

- Revenue
- Gross Profit (GP)
- Profit Before Tax (PBT)
- Profit After Tax (PAT)
- GP Margin
- PBT Margin
- PAT Margin

After extracting these values, compute YoY % for each metric.

------------------------------------------------------------
STRICT EXTRACTION RULES
------------------------------------------------------------
1. CLEANED JSON is the PRIMARY source. You MUST extract values from it first.
2. If a value is missing or blank in CLEANED JSON, then search RAW EXCEL TEXT and use that value. Must find all value.
3. If the value exists in neither source → return "blank".
4. You MUST include ALL years found in the data. Never skip a year.
5. All numeric output MUST:
   - be plain numbers (no commas, no RM)
   - convert percentages to numeric (e.g., "23%" → 23)
6. If the input contains "---", "-", "n/a", empty cell, or unreadable text → treat as blank.
7. YoY calculation rules:
   -Calculate ALL part of yoy for available data. 
   - If both years contain numeric values: YoY = (year2 - year1) / year1 * 100
   - If any value is blank or not numeric → YoY = "blank"
8. DO NOT guess any value. Only extract values that appear explicitly in either dataset.
9. Output must be STRICTLY VALID JSON. No text, no markdown, no comments.
 

------------------------------------------------------------
OUTPUT FORMAT (STRICT)
------------------------------------------------------------
Return ONLY valid JSON in this exact structure:

{
  "years": {
      "<year>": {
          "Revenue": "",
          "GP": "",
          "PBT": "",
          "PAT": "",
          "GP Margin": "",
          "PBT Margin": "",
          "PAT Margin": ""
      }
  },
  "yoy_change": {
      "Revenue": "",
      "GP": "",
      "PBT": "",
      "PAT": "",
      "GP Margin": "",
      "PBT Margin": "",
      "PAT Margin": ""
  }
}

------------------------------------------------------------
SOURCE DATA
------------------------------------------------------------

CLEANED_JSON:
""" + json.dumps(clean_json) + """


RAW_EXCEL_TEXT:
""" + combined_text

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Return only valid JSON. No markdown."},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )
    update_token_usage(response.usage)
    raw = response.choices[0].message.content
    if raw.startswith("```"):
        raw = raw.strip("```").replace("json", "").strip()

    return json.loads(raw)  # keep in memory

# ------------------------
# 5. Print & save financial table → now returns text
# ------------------------
def format_val(val, is_ratio=False, width=12, scale_factor=1):

    if val in ["blank", None, ""]:
        return str(val).rjust(width)

    # Try to convert to float
    try:
        num = float(val) / scale_factor
    except:
        return str(val).rjust(width)

    # Cap to maximum 6 digits for display (integer part)
    if not is_ratio:
        num = max(min(num, 999999), -999999)

    # Handle negative values with parentheses
    if num < 0:
        if is_ratio:
            num_str = f"({abs(num):.2f})"
        else:
            num_str = f"({abs(int(num)):,})"
    else:
        if is_ratio:
            num_str = f"{num:.2f}"
        else:
            num_str = f"{int(num):,}"

    # Right justify the whole string
    return num_str.rjust(width)


    # Integer formatting with max 6 digits
    try:
        num_int = int(num)
        # Handle negatives with parentheses
        if num_int < 0:
            num_str = f"({abs(num_int):,})"
        else:
            num_str = f"{num_int:,}"

        # Truncate if longer than width
        if len(num_str) > width:
            num_str = num_str[-width:]

        return num_str.rjust(width)
    except:
        return str(num).rjust(width)


def financial_report_text(data):
    years = list(data.get("years", {}).keys())
    if len(years) < 2:
        return "⚠️ Need at least 2 years of data for YoY calculations."
    year1, year2 = years[0], years[1]
    col_metric = 45; col_val=12; col_yoy=12
    lines = []
    lines.append("\n" + "Financial Result".ljust(col_metric) + f"FYE{year1}".rjust(col_val) + f"FYE{year2}".rjust(col_val) + "YoY changes".rjust(col_yoy))
    lines.append("_" * (col_metric + col_val*2 + col_yoy))
    
    # Indicators
    lines.append("Financial indicators (RM'000)")
    indicators = ["Revenue","GP","PBT","PAT"]
    for m in indicators:
        val1 = data["years"].get(year1, {}).get(m,"blank")
        val2 = data["years"].get(year2, {}).get(m,"blank")
        yoy = data["yoy_change"].get(m,"blank")
        
        # Divide numeric values by 1000
        val1 = float(val1)/1000 if val1 not in ["blank", None, ""] else val1
        val2 = float(val2)/1000 if val2 not in ["blank", None, ""] else val2

        lines.append(f"{m:<{col_metric}}{format_val(val1)}{format_val(val2)}{format_val(yoy)}")
    
    lines.append("_" * (col_metric + col_val*2 + col_yoy))
    
    # Ratios
    lines.append("Financial Ratios (%)")
    ratios = ["GP Margin","PBT Margin","PAT Margin"]
    for m in ratios:
        val1 = data["years"].get(year1, {}).get(m,"blank")
        val2 = data["years"].get(year2, {}).get(m,"blank")
        yoy = data["yoy_change"].get(m,"blank")
        lines.append(f"{m:<{col_metric}}{format_val(val1,is_ratio=True)}{format_val(val2,is_ratio=True)}{format_val(yoy,is_ratio=True)}")
    
    report_text = "\n".join(lines)
    return report_text

# ------------------------
# 6. Revenue JSON via GPT
# ------------------------
def ask_ai_for_revenue(raw_text_dict, clean_json):
    combined_text = "\n\n".join(f"{name}:\n{text}" for name, text in raw_text_dict.items())

    prompt = f"""
You are a financial analyst. Using BOTH raw Excel text and cleaned JSON, extract ONLY the following:

- Revenue by segment for ALL years appearing in the data.
- Include all segments exactly as they appear in Excel.
- Calculate YoY % for every segment where two consecutive years exist. (8% instead of 0.08)
- Detect the years automatically.

Return JSON in this format:

{{
  "years": {{
      "<year_1>": {{
          "Revenue": {{}}
      }},
      "<year_2>": {{
          "Revenue": {{}}
      }}
  }},
  "yoy_change": {{
      "Revenue": {{
          "<segment>": {{
              "<year_x_to_year_y>": "<percentage>"
          }}
      }}
  }}
}}

RULES:
1. Use cleaned JSON as primary.
2. Missing values → "blank".
3. YoY only if both years numeric.
4. Use whole numbers without commas.
5. Return ONLY JSON.

CLEANED JSON:
{json.dumps(clean_json)}

RAW EXCEL TEXT:
{combined_text}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Return only JSON."},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )
    update_token_usage(response.usage)
    raw = response.choices[0].message.content
    if raw.startswith("```"):
        raw = raw.strip("```").replace("json", "").strip()

    return json.loads(raw)  # keep in memory

# ------------------------
# 7. Revenue table → returns text
# ------------------------
def revenue_table_text(data):
    years = sorted(data["years"].keys(), key=lambda x:int(x))
    if len(years)<2:
        return "⚠️ Need at least 2 years."
    year1,year2 = years[0],years[1]
    revenue1 = data["years"][year1]["Revenue"]
    revenue2 = data["years"][year2]["Revenue"]
    yoy = data["yoy_change"].get("Revenue",{})
    col_segment = 30; col_val=18; col_yoy=18
    lines=[]
    lines.append(f"{'Segment':<{col_segment}}{'FYE'+year1:>{col_val}}{'FYE'+year2:>{col_val}}{'YoY (%)':>{col_yoy}}")
    lines.append("_"*(col_segment+col_val*2+col_yoy))
    all_segments = sorted(set(revenue1.keys())|set(revenue2.keys()))
    for seg in all_segments:
        val1= revenue1.get(seg,"blank"); val2=revenue2.get(seg,"blank")
        yoy_val = yoy.get(seg,{}).get(f"{year1}_to_{year2}","blank")
        val1_str = f"{int(val1):,}" if isinstance(val1,(int,float)) else "blank"
        val2_str = f"{int(val2):,}" if isinstance(val2,(int,float)) else "blank"
        if yoy_val!="blank":
            num=float(yoy_val); yoy_str=f"({abs(num):.2f}%)" if num<0 else f"{num:.2f}%"
        else: yoy_str="blank"
        lines.append(f"{seg:<{col_segment}}{val1_str:>{col_val}}{val2_str:>{col_val}}{yoy_str:>{col_yoy}}")
    report_text="\n".join(lines)
    return report_text

# ------------------------
# 8. Revenue narrative
# ------------------------
def generate_revenue_narrative(revenue_data, raw_text_dict):
    json_text = json.dumps(revenue_data,indent=2)
    raw_excel_text = "\n\n".join(f"{n}:\n{text}" for n,text in raw_text_dict.items())
    prompt=f"""
Prepare MD&A revenue section.

1. Paragraph 1: Total revenue movement, RM change, %, causes (infer from raw text)
2. Then one paragraph per segment: revenue year1, year2, RM change, YoY %, reason (use raw text) (No need title for each segment)

DATA:
{json_text}

RAW EXCEL:
{raw_excel_text}
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"system","content":"Write MD&A financial analysis."},{"role":"user","content":prompt}],
        temperature=0.2
    )
    update_token_usage(response.usage)
    text=response.choices[0].message.content.strip()
    return text  # keep in memory

# ------------------------
# 9. Donut Chart → returns BytesIO
# ------------------------
def generate_donut_chart(data):
    years=sorted(data["years"].keys(),key=lambda x:int(x))
    year1,year2 = years[0],years[1]
    rev1 = data["years"][year1]["Revenue"]
    rev2 = data["years"][year2]["Revenue"]

    # scale to RM (if original numbers in JSON are thousands)
    def scale(rev): return {k:v*1000 if isinstance(v,(int,float)) else v for k,v in rev.items()}
    rev1=scale(rev1); rev2=scale(rev2)

    # percentages
    def pct(rev):
        total=sum(v for v in rev.values() if isinstance(v,(int,float)))
        return {k:v/total*100 for k,v in rev.items()}, total
    pct1,_=pct(rev1); pct2,_=pct(rev2)

    fig,axes=plt.subplots(1,2,figsize=(14,6))
    def draw(ax,title,percentages):
        labels=list(percentages.keys())
        values=list(percentages.values())
        ax.pie(values, labels=[f"{v:.1f}%" for v in values], startangle=90, wedgeprops=dict(width=0.35),textprops={'fontsize': 16})
        ax.set_title(title, fontsize=16, pad=20)
        ax.legend(labels, loc="lower center", bbox_to_anchor=(0.5,-0.25))
    draw(axes[0], f"{year1} Revenue Contribution", pct1)
    draw(axes[1], f"{year2} Revenue Contribution", pct2)
    plt.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format="PNG", dpi=300, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf  # return in-memory image

# ------------------------
# 10. GP & PBT Commentary + Charts → returns text + BytesIO
# ------------------------
def generate_gp_pbt_analysis(data):
    raw_years = sorted(data["years"].keys(), key=lambda x: int(x))
    years = [f"FYE {y}" for y in raw_years]

    # Prepare GP & GP Margin
    gp_values = [float(data["years"][y].get("GP", 0)) for y in raw_years]
    gp_margin_values = [float(data["years"][y].get("GP Margin", 0)) for y in raw_years]
    prev_year = raw_years[-2]
    curr_year = raw_years[-1]
    gp_prev = float(data["years"][prev_year]["GP"])
    gp_curr = float(data["years"][curr_year]["GP"])
    gp_margin_prev = float(data["years"][prev_year]["GP Margin"])
    gp_margin_curr = float(data["years"][curr_year]["GP Margin"])

    # Prepare PBT & PBT Margin
    pbt_values = [float(data["years"][y].get("PBT", 0)) for y in raw_years]
    pbt_margin_values = [float(data["years"][y].get("PBT Margin", 0)) for y in raw_years]
    pbt_prev = float(data["years"][prev_year]["PBT"])
    pbt_curr = float(data["years"][curr_year]["PBT"])
    pbt_margin_prev = float(data["years"][prev_year]["PBT Margin"])
    pbt_margin_curr = float(data["years"][curr_year]["PBT Margin"])

    def generate_commentary(metric_name, prev_val, curr_val, prev_margin, curr_margin):
        prompt = f"""
Generate a financial analysis paragraph for an annual report.

Data:
- {metric_name} in FYE {prev_year}: RM{prev_val} million
- {metric_name} in FYE {curr_year}: RM{curr_val} million
- {metric_name} Margin in FYE {prev_year}: {prev_margin}%
- {metric_name} Margin in FYE {curr_year}: {curr_margin}%

Requirements:
1. Write 1 paragraph only.
2. Mention increase/decrease and include RM change & % change.
3. Also describe Margin increase/decrease.
4. Tone should match: “In line with the Group’s operational performance, the Group’s {metric_name} recorded an increase …”
5. Annual report style, formal, polished.

Output only the final paragraph.
"""
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        update_token_usage(response.usage)
        return response.choices[0].message.content.strip()

    gp_commentary = generate_commentary("GP", gp_prev, gp_curr, gp_margin_prev, gp_margin_curr)
    pbt_commentary = generate_commentary("PBT", pbt_prev, pbt_curr, pbt_margin_prev, pbt_margin_curr)

    def plot_metric_bar_line(metric_values, margin_values, title):
        fig, ax1 = plt.subplots(figsize=(7, 4))

        # ---- Chart Title ----
        fig.suptitle(title, fontsize=14, fontweight='medium', y=1.02)

        # ---- Bar Chart ----
        bars = ax1.bar(years, metric_values, color='skyblue')
        for bar, value in zip(bars, metric_values):
            ax1.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() * 0.5,
                f"RM{value} mil",
                ha='center',
                va='center',
                fontsize=10,
                color='white'
            )

        # ---- Line Chart ----
        ax2 = ax1.twinx()
        ax2.plot(years, margin_values, marker='o', color='red')

        for x, y in zip(years, margin_values):
            ax2.text(x, y + 0.05, f"{y}%", ha='center', fontsize=10)

        # Remove ticks & spines for a clean look
        ax1.set_yticks([]); ax2.set_yticks([])
        for spine in ax1.spines.values(): spine.set_visible(False)
        for spine in ax2.spines.values(): spine.set_visible(False)
        ax1.tick_params(left=False, bottom=False)

        plt.tight_layout()

        buf = BytesIO()
        fig.savefig(buf, format="PNG", dpi=300, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    gp_chart = plot_metric_bar_line(gp_values, gp_margin_values, "GP and GP Margin")
    pbt_chart = plot_metric_bar_line(pbt_values, pbt_margin_values, "PBT and PBT Margin")


    return gp_commentary, pbt_commentary, gp_chart, pbt_chart

# ------------------------
# Word export → returns in-memory BytesIO
# ------------------------
def save_all_outputs_to_word(
        business_text, financial_text, revenue_text,
        revenue_narrative, gp_commentary, pbt_commentary,
        revenue_chart_buf, gp_chart_buf, pbt_chart_buf
    ):

    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()

    def add_paragraph_with_spacing(text, space_after=Pt(12)):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = space_after

    # --------------------------------------------------
    # TEXT SECTIONS (title_flag = True)
    # --------------------------------------------------
    text_sections_ordered = [
        ("Business & Operational Overview", business_text, True),
        ("Review of financial performance", financial_text, True),
        ("Revenue Table", revenue_text, True),
        ("Revenue Narrative", revenue_narrative, True),
        ("GP and GP Margin", gp_commentary, True),
        ("PBT and PBT Margin", pbt_commentary, True),
    ]

    # --------------------------------------------------
    # IMAGE SECTIONS (title_flag = False)
    # --------------------------------------------------
    image_sections_ordered = [
        ("Revenue Contribution (Donut)", revenue_chart_buf, False),
        ("Gross Profit Chart (GP)",         gp_chart_buf, False),
        ("PBT Chart (PBT)",                 pbt_chart_buf, False),
    ]

    # --------------------------------------------------
    # 1. Business Overview + Financial Performance
    # --------------------------------------------------
    for title, text, show_title in text_sections_ordered[:2]:
        if show_title:
            doc.add_heading(title, level=1)
        for line in text.split("\n"):
            add_paragraph_with_spacing(line)

    # --------------------------------------------------
    # 2. Page break before Revenue Table
    # --------------------------------------------------
    doc.add_page_break()

    # --------------------------------------------------
    # 3. Revenue Table
    # --------------------------------------------------
    title, text, show_title = text_sections_ordered[2]
    if show_title:
        doc.add_heading(title, level=1)
    for line in text.split("\n"):
        add_paragraph_with_spacing(line)

    # --------------------------------------------------
    # 4. Revenue Chart (NO TITLE)
    # --------------------------------------------------
    _, buf, show_title = image_sections_ordered[0]
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(buf, width=Inches(5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_with_spacing("")

    # --------------------------------------------------
    # 5. Revenue Narrative (with title)
    # --------------------------------------------------
    title, text, show_title = text_sections_ordered[3]
    if show_title:
        doc.add_heading(title, level=1)
    for line in text.split("\n"):
        add_paragraph_with_spacing(line)

    # --------------------------------------------------
    # 6. GP and GP Margin (new page)
    # --------------------------------------------------
    doc.add_page_break()
    title, text, show_title = text_sections_ordered[4]
    if show_title:
        doc.add_heading(title, level=1)
    for line in text.split("\n"):
        add_paragraph_with_spacing(line)

    # --------------------------------------------------
    # 7. GP Chart (NO TITLE)
    # --------------------------------------------------
    _, buf, show_title = image_sections_ordered[1]
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(buf, width=Inches(5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_with_spacing("")

    # --------------------------------------------------
    # 8. PBT and PBT Margin
    # --------------------------------------------------
    title, text, show_title = text_sections_ordered[5]
    if show_title:
        doc.add_heading(title, level=1)
    for line in text.split("\n"):
        add_paragraph_with_spacing(line)

    # --------------------------------------------------
    # 9. PBT Chart (NO TITLE)
    # --------------------------------------------------
    _, buf, show_title = image_sections_ordered[2]
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(buf, width=Inches(5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_with_spacing("")

    # --------------------------------------------------
    # Return as BytesIO
    # --------------------------------------------------
    out_buf = BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)
    return out_buf



total_input_tokens = 0
total_output_tokens = 0
PRICE_INPUT = 0.00000015
PRICE_OUTPUT = 0.00000060
USD_TO_RM = 4.14

# ------------------------
# Token usage updater
# ------------------------
def update_token_usage(usage):
    """Update global counters and print token/cost summary."""
    global total_input_tokens, total_output_tokens
    input_tokens = getattr(usage, "prompt_tokens", 0)
    output_tokens = getattr(usage, "completion_tokens", 0)

    total_input_tokens += input_tokens
    total_output_tokens += output_tokens

    cost_input = total_input_tokens * PRICE_INPUT
    cost_output = total_output_tokens * PRICE_OUTPUT
    total_cost_usd = cost_input + cost_output
    total_cost_rm = total_cost_usd * USD_TO_RM

    print("\n=========== GPT TOKEN USAGE FOR MD&A================")
    print(f"Last call input tokens:  {input_tokens}")
    print(f"Last call output tokens: {output_tokens}")
    print(f"Total input tokens:  {total_input_tokens:,}")
    print(f"Total output tokens: {total_output_tokens:,}")
    print(f"Cost for input (USD):  ${cost_input:.8f}")
    print(f"Cost for output (USD): ${cost_output:.8f}")
    print(f"-----------------------------------------------------")
    print(f"Total cost (USD):            ${total_cost_usd:.8f}")
    print(f"Estimated Total cost (RM):   RM {total_cost_rm:.4f}")
    print("======================================================\n")

# ------------------------
# MAIN PIPELINE → returns all outputs in memory
# ------------------------
if __name__ == "__main__":
    # Example: uploaded_files_dict = {"MD&A.xlsx": "path_or_fileobj", "consol.xlsx": "path_or_fileobj"}
    # Replace with your actual file upload logic
    uploaded_files_dict = {}  # ← populate with uploaded files

    if "MD&A.xlsx" not in uploaded_files_dict:
        raise Exception("MD&A.xlsx not uploaded.")
    if "consol.xlsx" not in uploaded_files_dict:
        raise Exception("consol.xlsx not uploaded.")

    mdna_file_obj = uploaded_files_dict.get("MD&A.xlsx")
    consol_file_obj = uploaded_files_dict.get("consol.xlsx")

    # Load Excel content
    mdna_text = load_excel_clean(mdna_file_obj)
    consol_text = load_excel_clean(consol_file_obj)

    # Business overview
    business_text = generate_business_overview(mdna_text, consol_text)

    # Only consol.xlsx can generate clean_json
    clean_json = load_consol_to_json(consol_file_obj, "GRP IS")

    # Combine all sheets
    all_sheets_text = {**mdna_text, **consol_text}

    # Financial data
    financial_data = ask_ai_for_financials(all_sheets_text, clean_json)
    financial_text = financial_report_text(financial_data)

    # Revenue
    revenue_data = ask_ai_for_revenue(all_sheets_text, clean_json)
    revenue_text = revenue_table_text(revenue_data)
    revenue_narrative = generate_revenue_narrative(revenue_data, all_sheets_text)
    revenue_chart_buf = generate_donut_chart(revenue_data)

    # GP & PBT analysis
    gp_commentary, pbt_commentary, gp_chart_buf, pbt_chart_buf = generate_gp_pbt_analysis(financial_data)

    # Generate Word report in memory
    word_buf = save_all_outputs_to_word(
        business_text, financial_text, revenue_text,
        revenue_narrative, gp_commentary, pbt_commentary,
        revenue_chart_buf, gp_chart_buf, pbt_chart_buf
    )

    # Save Word report to disk (now word_buf is guaranteed to exist)
    output_path = "Annual_Report.docx"
    with open(output_path, "wb") as f:
        f.write(word_buf.getbuffer())

    print(f"Word report saved to {output_path}")




  
